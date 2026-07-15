import fitz
import pdfplumber
import docx
import os
import base64
import requests
import io
from PIL import Image
from dotenv import load_dotenv
load_dotenv()

HF_TOKEN = os.environ.get("HF_TOKEN", "")

HF_API_URL = "https://api-inference.huggingface.co/models/llava-hf/llava-1.5-7b-hf"

def get_hf_token():
    """Read token fresh each time — ensures dotenv is loaded first."""
    token = os.environ.get("HF_TOKEN", "")
    print(f"[Parser] HF_TOKEN loaded: {bool(token)}")
    print(f"[Parser] HF_TOKEN prefix: {token[:8] if token else 'EMPTY'}")
    return token

def image_to_base64(image: Image.Image) -> str:
    """Convert PIL Image to base64 string."""
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return base64.b64encode(buffer.getvalue()).decode("utf-8")


def analyze_image_with_groq_vision(image: Image.Image) -> str:
    """Use Groq's vision model — free and reliable."""
    try:
        from groq import Groq
        client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

        # Convert to base64
        buffer = io.BytesIO()
        image.save(buffer, format="PNG")
        img_b64 = base64.b64encode(buffer.getvalue()).decode("utf-8")

        response = client.chat.completions.create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{img_b64}"
                            }
                        },
                        {
                            "type": "text",
                            "text": """Extract ALL content from this document image:
- All text exactly as written
- Table contents with structure
- Chart/graph descriptions
- Diagram descriptions
- Handwritten text
- Headers and footers
Return everything in structured format."""
                        }
                    ]
                }
            ],
            max_tokens=1500,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"[Parser] Groq vision error: {e}")
        return ""


def extract_text_from_pdf(file_path: str) -> dict:
    """
    Extract text from PDF.
    For pages with little/no text → use LLaVA vision.
    For normal text pages → use PyMuPDF directly.
    """
    doc = fitz.open(file_path)
    pages_content = []
    total_pages = len(doc)
    used_vision = False

    for page_num in range(total_pages):
        page = doc[page_num]
        text = page.get_text().strip()

        # Check if page has images or very little text
        has_images = len(page.get_images()) > 0
        is_scanned = len(text) < 100

        if (is_scanned or has_images) and HF_TOKEN:
            # Use LLaVA vision for this page
            try:
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                pil_image = Image.open(io.BytesIO(img_data))

                vision_text = analyze_image_with_llava(pil_image)

                if vision_text and len(vision_text) > len(text):
                    text = vision_text
                    used_vision = True
                    print(f"[Parser] Used vision for page {page_num + 1}")
            except Exception as e:
                print(f"[Parser] Vision failed for page {page_num + 1}: {e}")
                # Fallback to pdfplumber
                try:
                    with pdfplumber.open(file_path) as pdf:
                        if page_num < len(pdf.pages):
                            pb_text = pdf.pages[page_num].extract_text()
                            if pb_text:
                                text = pb_text
                except Exception:
                    pass

        pages_content.append(f"\n--- Page {page_num + 1} ---\n{text}")

    doc.close()

    return {
        "text": "\n".join(pages_content).strip(),
        "total_pages": total_pages,
        "used_vision": used_vision
    }


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from Word documents."""
    try:
        doc = docx.Document(file_path)
        sections = []
        for para in doc.paragraphs:
            if para.text.strip():
                sections.append(para.text)
        for table in doc.tables:
            table_text = "[TABLE]\n"
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                )
                if row_text.strip():
                    table_text += row_text + "\n"
            sections.append(table_text)
        return "\n".join(sections)
    except Exception as e:
        print(f"[Parser] DOCX error: {e}")
        return ""


def extract_text_from_image(file_path: str) -> str:
    """
    Extract content from image files using LLaVA.
    Truly multimodal — understands charts, diagrams etc.
    """
    try:
        pil_image = Image.open(file_path)

        result = analyze_image_with_groq_vision(pil_image)
        if result:
           return result

        # Fallback to pytesseract if available
        try:
            import pytesseract
            return pytesseract.image_to_string(pil_image)
        except ImportError:
            pass

        return f"[Image file: {os.path.basename(file_path)}]"

    except Exception as e:
        print(f"[Parser] Image error: {e}")
        return ""


def extract_text(file_path: str) -> dict:
    """Main extraction — auto-detects file type."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)

    elif ext in [".docx", ".doc"]:
        text = extract_text_from_docx(file_path)
        return {"text": text, "total_pages": 1, "used_vision": False}

    elif ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()
        return {"text": text, "total_pages": 1, "used_vision": False}

    elif ext in [".jpg", ".jpeg", ".png", ".webp"]:
        text = extract_text_from_image(file_path)
        return {"text": text, "total_pages": 1, "used_vision": True}

    return {"text": "", "total_pages": 0, "used_vision": False}


def get_document_metadata(file_path: str) -> dict:
    """Get full document metadata."""
    result = extract_text(file_path)
    text = result.get("text", "")
    word_count = len(text.split())
    estimated_tokens = int(word_count / 0.75)

    return {
        "text": text,
        "word_count": word_count,
        "char_count": len(text),
        "estimated_tokens": estimated_tokens,
        "total_pages": result.get("total_pages", 1),
        "used_vision": result.get("used_vision", False),
        "needs_rag": estimated_tokens > 100000
    }